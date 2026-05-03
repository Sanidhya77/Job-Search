"""CV reader tool.

Reads the user's CV from disk and returns plain text. Supports PDF
(via pdfplumber) and DOCX (via python-docx). The format is dispatched
on the file extension because reading the magic bytes would add a
dependency without solving a problem the user actually has - they
always know what file they uploaded.

This is a deterministic tool, not an LLM step. The LLM-based CV
analyser later turns the plain text into a structured UserProfile.
Keeping the two steps separate means the file format never reaches
the LLM, which keeps prompts portable across PDF and DOCX inputs.
"""

from pathlib import Path

import pdfplumber
from docx import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class CVReaderError(Exception):
    """Raised when the CV cannot be read for any reason.

    A single exception type for the whole module means callers do
    not need to catch a different exception per format. The message
    carries the specific cause.
    """


def read_cv(path: Path) -> str:
    """Read a CV file and return its text content.

    Args:
        path: Path to a .pdf or .docx file.

    Returns:
        The full text of the CV with normalised whitespace.

    Raises:
        CVReaderError: If the file is missing, has an unsupported
            extension, or cannot be parsed.
    """
    path = Path(path)

    if not path.exists():
        raise CVReaderError(f"CV file not found: {path}")

    if not path.is_file():
        raise CVReaderError(f"CV path is not a file: {path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise CVReaderError(
            f"Unsupported CV format {suffix!r}. "
            f"Supported formats: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    try:
        if suffix == ".pdf":
            text = _read_pdf(path)
        else:
            text = _read_docx(path)
    except CVReaderError:
        raise
    except Exception as exc:
        # Wrap any underlying parser error in our own type so callers
        # only have to handle CVReaderError.
        raise CVReaderError(f"Failed to parse {path.name}: {exc}") from exc

    normalised = _normalise_whitespace(text)
    if not normalised:
        raise CVReaderError(f"CV file {path.name} contains no extractable text.")

    return normalised


def _read_pdf(path: Path) -> str:
    """Extract text from every page of a PDF and join with newlines."""
    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text)
    return "\n".join(pages)


def _read_docx(path: Path) -> str:
    """Extract text from every paragraph of a DOCX file.

    Tables are walked separately because python-docx does not include
    table text in document.paragraphs. Most CVs put contact info or
    skills sections in tables, so missing them would lose real data.
    """
    doc = Document(str(path))

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return "\n".join(parts)


def _normalise_whitespace(text: str) -> str:
    """Collapse runs of blank lines and trim trailing spaces.

    CV PDFs often contain layout-driven extra blank lines that confuse
    downstream LLM parsing without adding information. We keep single
    blank lines as paragraph separators but collapse longer runs.
    """
    lines = [line.rstrip() for line in text.splitlines()]

    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if line.strip() == "":
            blank_run += 1
            if blank_run <= 1:
                cleaned.append("")
        else:
            blank_run = 0
            cleaned.append(line)

    return "\n".join(cleaned).strip()
